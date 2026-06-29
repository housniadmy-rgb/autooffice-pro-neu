"""Build PraxisOnline24_Investment_Report.docx — professional M&A-style report."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent / "PraxisOnline24_Investment_Report.docx"

doc = Document()

# ── Default styles ───────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

def add_heading(text, level=1, color=(31, 73, 125)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def add_kv_table(rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Light Grid Accent 1'
    for i, (k, v) in enumerate(rows):
        t.cell(i, 0).text = k
        t.cell(i, 1).text = v
        for cell in t.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
        # bold the key column
        for r in t.cell(i, 0).paragraphs[0].runs:
            r.bold = True
    return t

# ── Title page ───────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("PraxisOnline24")
r.font.size = Pt(36)
r.bold = True
r.font.color.rgb = RGBColor(31, 73, 125)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Investment & Due-Diligence Report")
r.font.size = Pt(18)
r.italic = True
r.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
slogan = doc.add_paragraph()
slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = slogan.add_run("Ihre Praxis. Online. Rund um die Uhr.")
r.font.size = Pt(14)
r.italic = True

doc.add_paragraph()
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Pre-Revenue SaaS Asset for Acquisition\nProduct: Online appointment, practice management & multi-language booking platform")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(110, 110, 110)

doc.add_paragraph()
url = doc.add_paragraph()
url.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = url.add_run("https://www.praxisonline24.com")
r.font.size = Pt(13)
r.bold = True

doc.add_page_break()

# ── Table of contents notice ─────────────────────────────────────────────
add_heading("Contents", level=1)
toc = [
    "1.  Executive Summary",
    "2.  Problem",
    "3.  Solution",
    "4.  Target Audience",
    "5.  Market",
    "6.  Product Overview",
    "7.  Architecture",
    "8.  Technologies",
    "9.  Features",
    "10. Database",
    "11. APIs",
    "12. Security",
    "13. AI Capabilities",
    "14. Deployment",
    "15. SEO",
    "16. Roadmap",
    "17. Project Statistics",
    "18. Opportunities",
    "19. Risks",
    "20. Market Potential",
    "21. Valuation",
    "22. Sales Arguments",
]
for line in toc:
    add_para(line, size=11)

doc.add_page_break()

# ── 1. Executive Summary ─────────────────────────────────────────────────
add_heading("1. Executive Summary", level=1)
add_para(
    "PraxisOnline24 is a fully built, production-deployed B2B SaaS for medical practices, therapists, "
    "and small service businesses. The product covers the entire patient lifecycle — from a public "
    "24/7 online booking flow through admin calendar, invoicing (PDF), waitlist auto-offer, review "
    "moderation, to a rule-based 'AI Practice Manager' dashboard. The codebase ships in 12 languages "
    "(including RTL Arabic) and runs live at https://www.praxisonline24.com on Render.com Frankfurt, "
    "fronted by Cloudflare.", size=11)
add_para("")
add_para("Status:", bold=True)
add_bullet("16 development phases shipped (see README.md)")
add_bullet("Live production deployment with verified TLS")
add_bullet("12-language i18n (DE/EN/FR/ES/PT/AR/RU/ZH/HI/TH/TR/ID)")
add_bullet("Privacy-by-design data model: no health data, automatic anonymization after 30 days")
add_bullet("Pre-revenue: zero customers, zero MRR, zero ARR")
add_para("")
add_para("Investment Thesis:", bold=True)
add_para(
    "Buy a production-ready, multi-language SaaS asset at a fraction of replacement cost (~€75k–€120k "
    "for a from-scratch rebuild). Ideal for operators with go-to-market capacity in EU, DACH, or "
    "any of the 12 supported language markets. The 'AI Practice Manager' module + 12-language coverage "
    "are the two strongest differentiators in the price band.")

doc.add_page_break()

# ── 2. Problem ───────────────────────────────────────────────────────────
add_heading("2. Problem", level=1)
add_para(
    "Small medical practices and similar appointment-based service businesses face four chronic "
    "operational problems:")
add_bullet("Patients can only book during office hours by phone, creating friction and lost appointments.")
add_bullet("No-shows and cancellations are not automatically refilled, wasting practitioner time.")
add_bullet("Generic tools (Calendly + spreadsheets) do not handle multi-practitioner availability, "
           "invoicing, waitlists, or review collection in one place.")
add_bullet("Established players (Doctolib, Jameda, samedi, Dr. Flex in DACH) are expensive and "
           "force practices into a marketplace they do not control.")
add_para("")
add_para(
    "The market under-serves the practice that wants a low-cost, self-hostable, brand-controlled "
    "tool — and serves international markets even less, where multi-language SaaS in this niche "
    "is rare.")

# ── 3. Solution ──────────────────────────────────────────────────────────
add_heading("3. Solution", level=1)
add_para(
    "PraxisOnline24 is an integrated, white-label-friendly platform that covers booking, calendar, "
    "patient management, practitioner availability, waitlist auto-offer, invoicing with PDF "
    "generation, review moderation, and a 14-section daily KI-driven dashboard.")
add_para("")
add_para("How it differs from the market:", bold=True)
add_bullet("12 languages (DE/EN/FR/ES/PT/AR with RTL/RU/ZH/HI/TH/TR/ID) — unusual at this price band.")
add_bullet("Self-hostable monolith — runs on a €7/month Render Starter plan.")
add_bullet("Privacy-by-design — no health data ever stored; 30-day anonymisation cron.")
add_bullet("Rule-based AI Practice Manager — 14 dashboards generated from operational data.")
add_bullet("Multi-tenant from day one — every business table carries a practice_id foreign key.")

doc.add_page_break()

# ── 4. Target Audience ───────────────────────────────────────────────────
add_heading("4. Target Audience", level=1)
add_kv_table([
    ("Primary persona", "Owner-operated practice, 3–10 practitioners, >50 appointments/week"),
    ("Tools they replace", "Phone bookings, paper calendar, Calendly + spreadsheets"),
    ("Verticals (same product)", "Doctors, therapists, physiotherapy, veterinary, beauty/wellness, "
                                 "driving schools, tax / legal consulting, coaching"),
    ("Budget profile", "€300–€1,500/year for software"),
    ("Privacy posture", "GDPR-sensitive; values 'no health data stored'"),
    ("Geographic fit", "DACH primary; EU + emerging markets via 12-language coverage"),
])

# ── 5. Market ────────────────────────────────────────────────────────────
add_heading("5. Market", level=1)
add_para("Addressable market signals:", bold=True)
add_bullet("Germany: ~204,000 doctors' practices (KBV registry, public).")
add_bullet("EU: >1 million medical and paramedical practices.")
add_bullet("Adjacent verticals (physio, beauty, coaches) multiply the addressable base 3–5×.")
add_bullet("Multilingual coverage opens markets with weaker incumbents (Turkey, Indonesia, "
           "MENA region via Arabic).")
add_para("")
add_para("Competitive landscape:", bold=True)
add_bullet("Doctolib — DACH market leader, marketplace play, expensive for the practice.")
add_bullet("Jameda — review-focused; appointment booking is secondary.")
add_bullet("samedi — feature-rich, enterprise-leaning pricing.")
add_bullet("Dr. Flex — modern UX, single-language focus.")
add_bullet("Calendly + manual workflows — low cost, missing invoicing / waitlist / multi-practitioner.")
add_para("")
add_para(
    "PraxisOnline24 differentiates through (a) price, (b) self-hostability, (c) 12-language coverage, "
    "(d) brand control. It does not compete with marketplace-driven players on lead-gen.")

# ── 6. Product Overview ──────────────────────────────────────────────────
add_heading("6. Product Overview", level=1)
add_para(
    "The product is built around three user perspectives: patient (public, no auth), practice staff "
    "(authenticated admin), and platform owner (root admin).")

add_para("Patient surface", bold=True)
add_bullet("3-step public booking flow: practitioner → date → time → details.")
add_bullet("ICS calendar download after booking.")
add_bullet("Cancel via one-time token e-mail link.")
add_bullet("Patient portal for appointment lookup by e-mail.")
add_bullet("Public review submission with disclaimer (no health data).")
add_bullet("Waitlist signup; auto-offer when a slot frees.")

add_para("Admin surface", bold=True)
add_bullet("Dashboard with metrics, today's priorities, AI Practice Manager insights.")
add_bullet("Calendar (day/week/month) with conflict detection and archive filter.")
add_bullet("Practitioner CRUD + per-practitioner weekly availability.")
add_bullet("Patient list derived from appointment history (privacy-by-design).")
add_bullet("Invoices with line items, VAT, PDF download.")
add_bullet("Review moderation (approve / hide / delete).")
add_bullet("Waitlist administration with auto-offer status.")
add_bullet("Recipe print log (audit trail without prescription content).")
add_bullet("Settings: profile, language, archive policy, account status, password change.")
add_bullet("Activity log (every admin action recorded).")

doc.add_page_break()

# ── 7. Architecture ──────────────────────────────────────────────────────
add_heading("7. Architecture", level=1)
add_para("High-level:", bold=True)
add_para("Cloudflare (TLS + CDN)  →  Render.com Web Service (Frankfurt)  →  Express monolith  →  SQLite via sql.js")
add_para("")
add_kv_table([
    ("Process model", "Single Node process, modular Express routes, in-memory session store"),
    ("Static assets", "Served by express.static before session/CORS — minimal middleware on cold paths"),
    ("Auth model", "Session-cookie + role-based middleware (admin > doctor > staff > patient)"),
    ("Auto-pause", "402 response if practice account is paused (post-trial)"),
    ("Cron", "node-cron in same process, 8 jobs"),
    ("Background safety", "Daily SQLite backup + activity log + automation_log"),
    ("Schema migration", "Lightweight CREATE-IF-NOT-EXISTS + ALTER-TABLE-add-column with try/catch"),
])

# ── 8. Technologies ──────────────────────────────────────────────────────
add_heading("8. Technologies", level=1)
add_kv_table([
    ("Runtime", "Node.js ≥ 18 LTS"),
    ("Server framework", "Express 4.18"),
    ("Database", "SQLite via sql.js 1.12 (single-file, easy backup)"),
    ("Auth", "express-session + bcrypt (cost 12)"),
    ("Frontend", "Plain HTML5 + CSS3 + vanilla ES modules (no framework lock-in)"),
    ("PDF", "pdfkit"),
    ("Scheduling", "node-cron"),
    ("E-mail", "nodemailer (SMTP) + Brevo API"),
    ("i18n", "Custom JSON catalogs (12 languages)"),
    ("Hosting", "Render.com (Frankfurt), render.yaml committed"),
    ("Domain & TLS", "praxisonline24.com via Cloudflare → Render origin"),
    ("Backups", "Daily SQLite copy, 7-day rolling retention"),
])

# ── 9. Features ──────────────────────────────────────────────────────────
add_heading("9. Features", level=1)
add_para("All 16 shipped phases (per README.md):", bold=True)
phases = [
    ("Phase 1", "Express server, SQLite, sessions"),
    ("Phase 2", "Registration, login, logout, bcrypt, 30-day trial"),
    ("Phase 3", "Calendar (month/day), appointment management"),
    ("Phase 4", "Practitioners, patients (derived), BASIC limits"),
    ("Phase 5", "Invoices (PDF), recipe print log"),
    ("Phase 6", "Reviews (submit, approve, hide), settings"),
    ("Phase 7", "Waitlist + cron automations (reminders, cleanup)"),
    ("Phase 8", "3-step booking flow, patient portal, ICS export"),
    ("Phase 9", "Multilanguage (8 langs), BASIC/UNLIMITED packages, trial lifecycle"),
    ("Phase 10", "Booking language switcher, public multilingual pages, stabilization"),
    ("Phase 11", "Full multilingual public pages (12 langs)"),
    ("Phase 12", "Multilingual AGB/Datenschutz, expanded hero"),
    ("Phase 13", "AI Practice Manager — auto practice summary & hints"),
    ("Phase 14", "Dashboard multilingual (10 pages × 12 langs, cookie sync at login)"),
    ("Phase 15", "i18n completeness — all missing keys for 12 languages"),
    ("Phase 16", "Appointment archival — auto archive, configurable threshold, calendar filter"),
]
add_kv_table(phases)

doc.add_page_break()

# ── 10. Database ─────────────────────────────────────────────────────────
add_heading("10. Database", level=1)
add_para(
    "SQLite via sql.js — single file at data/praxisonline24.db. Loaded into RAM on boot, "
    "synchronously flushed to disk on every write.")
add_para("")
add_para("18 tables:", bold=True)
tables = [
    "practices", "users", "practitioners", "appointments", "waitlist",
    "waitlist_offers", "reviews", "invoices", "recipe_print_logs", "payments",
    "settings", "automation_log", "password_reset_tokens",
    "practitioner_availability", "activity_log", "demo_requests",
    "invite_tokens",
]
for t in tables:
    add_bullet(t)
add_para("")
add_para("Privacy-by-design — NOT stored:", bold=True)
add_bullet("No diagnoses")
add_bullet("No medications / dosages")
add_bullet("No date of birth")
add_bullet("No insurance numbers")
add_bullet("No prescription content (only print log: who, for whom, when)")

# ── 11. APIs ─────────────────────────────────────────────────────────────
add_heading("11. APIs", level=1)
add_para("17 Express route modules under /api:", bold=True)
apis = [
    ("/api/auth", "Login, logout, register, forgot/reset password, password change"),
    ("/api/dashboard", "Aggregated KPIs for admin home"),
    ("/api/practices", "Practice settings, language, subscription"),
    ("/api/appointments", "Calendar CRUD + status workflow"),
    ("/api/practitioners", "Practitioner CRUD + availability windows"),
    ("/api/patients", "Patient list derived from appointments"),
    ("/api/invoices", "Invoice CRUD + PDF generation"),
    ("/api/recipes", "Recipe print log"),
    ("/api/reviews", "Review moderation"),
    ("/api/waitlist", "Waitlist CRUD + auto-offer"),
    ("/api/payments", "Payment recording"),
    ("/api/public", "Public booking (no auth)"),
    ("/api/demo", "Demo request inbox"),
    ("/api/owner", "Owner-level admin"),
    ("/api/activity-log", "Audit log access"),
    ("/api/ai-praxismanager", "Rule-based dashboard (14 sections)"),
    ("/api/health", "Liveness probe"),
]
add_kv_table(apis)

# ── 12. Security ─────────────────────────────────────────────────────────
add_heading("12. Security", level=1)
add_para("Implemented controls:", bold=True)
add_bullet("Password hashing: bcrypt cost factor 12")
add_bullet("Session cookies: HttpOnly always, Secure in production")
add_bullet("HSTS (max-age=31536000; includeSubDomains; preload) in production")
add_bullet("X-Frame-Options: SAMEORIGIN")
add_bullet("X-Content-Type-Options: nosniff")
add_bullet("Referrer-Policy: strict-origin-when-cross-origin")
add_bullet("Permissions-Policy: camera=(), microphone=(), geolocation=()")
add_bullet("Rate limiting: 5/min login, 3/min register, 3/min forgot-password")
add_bullet("Role-based access (admin > doctor > staff > patient)")
add_bullet("Account-paused 402 response gate (per-route)")
add_bullet("Parametrised SQL throughout (no string concatenation)")
add_bullet("Trust proxy + Render TLS + Cloudflare in front")
add_bullet("Activity logging on key admin actions")
add_bullet("One-time password reset tokens with expiry")
add_para("")
add_para("Not yet implemented (typical MVP gaps; documented in Technical Due Diligence):", bold=True)
add_bullet("Content-Security-Policy header")
add_bullet("CSRF tokens (same-origin sessions are the current defence)")
add_bullet("2FA for admin users")
add_bullet("E-mail verification at registration")

doc.add_page_break()

# ── 13. AI Capabilities ──────────────────────────────────────────────────
add_heading("13. AI Capabilities", level=1)
add_para(
    "The 'AI Practice Manager' module (routes/ai-praxismanager.js, 731 lines) is a deterministic "
    "rule engine — not an LLM. It consumes the practice's operational state (appointments, "
    "invoices, reviews, waitlist) and produces a 14-section JSON payload that the admin dashboard "
    "renders.")
add_para("Sections produced:", bold=True)
add_bullet("Daily report (appointments today / tomorrow / new patients / open invoices)")
add_bullet("Financial overview (revenue today / month / open / paid / overdue)")
add_bullet("Appointment analytics (top type / favorite practitioner / cancellations / utilization)")
add_bullet("Review analytics (average / positive / negative / unresolved)")
add_bullet("Waitlist analytics (waiting / oldest / possible slots / open offers)")
add_bullet("Localized hints (12 languages, generated per practice state)")
add_bullet("Today's priorities (critical → urgent → info)")
add_bullet("KI recommendations (impact-tagged)")
add_bullet("Revenue forecast (current month projected from 3-month average)")
add_bullet("7-day utilization forecast")
add_bullet("Warnings (red / yellow / green)")
add_bullet("Auto-todos")
add_bullet("Auto setup-todos (e.g., 'add practitioner', 'add availability', 'complete profile')")
add_bullet("Traffic-light recommendations (Finance / Utilization / Reviews / Waitlist)")
add_para("")
add_para(
    "Note for transparency: The label 'AI' is marketing in the product name. The implementation is "
    "rule-based and predictable. A real LLM (Anthropic Claude or OpenAI) can be added by feeding "
    "the JSON payload into a summarisation prompt — approximately 1–2 days of integration work.")

# ── 14. Deployment ───────────────────────────────────────────────────────
add_heading("14. Deployment", level=1)
add_kv_table([
    ("Live URL", "https://www.praxisonline24.com"),
    ("Hosting", "Render.com Free Tier (currently); upgrade path to Starter = $7/mo"),
    ("Region", "Frankfurt (EU central)"),
    ("Front", "Cloudflare (TLS + CDN)"),
    ("Build", "npm install"),
    ("Start", "npm start (node server.js)"),
    ("Health", "GET /api/health → {\"status\":\"ok\"}"),
    ("Config", "render.yaml committed; one-click redeploy under buyer account"),
    ("Backups", "Daily SQLite copy in data/backups/, 7-day rolling"),
    ("Auto-deploy", "GitHub main branch → Render auto-deploys"),
])

# ── 15. SEO ──────────────────────────────────────────────────────────────
add_heading("15. SEO", level=1)
add_para("On-page SEO already implemented:", bold=True)
add_bullet("Per-page <title> with primary keyword")
add_bullet("<meta name='description'> on homepage + pricing")
add_bullet("Canonical URL")
add_bullet("OpenGraph + Twitter Card")
add_bullet("JSON-LD Schema.org SoftwareApplication on homepage")
add_bullet("sitemap.xml + robots.txt in public/")
add_para("")
add_para("Outstanding SEO work (post-acquisition wins):", bold=True)
add_bullet("hreflang tags for 12 languages (~4 hours)")
add_bullet("Lighthouse audit + image lazy-loading")
add_bullet("Backlink building (post-rebrand)")

doc.add_page_break()

# ── 16. Roadmap ──────────────────────────────────────────────────────────
add_heading("16. Roadmap (Buyer-Driven, Recommended)", level=1)
add_para("First 30 days:", bold=True)
add_bullet("Wipe test data, rotate secrets, set up own Render account")
add_bullet("Add CSP header, CSRF, 2FA for admins")
add_bullet("Re-enable Stripe checkout")
add_bullet("Add Sentry / error tracking")
add_bullet("Lawyer review of AGB / Datenschutz / Impressum")
add_para("First 90 days:", bold=True)
add_bullet("Migrate to PostgreSQL (~20–40h)")
add_bullet("Replace 'moment' with 'date-fns' (~1 day)")
add_bullet("Write ~30 Playwright E2E specs")
add_bullet("Native review of RU/ZH/HI/TH translations")
add_bullet("Launch paid acquisition campaigns in 1–2 verticals")
add_para("First 12 months:", bold=True)
add_bullet("First 50–200 paying practices via direct outreach + content marketing")
add_bullet("Vertical-specific UI variants (physio / beauty / coaching)")
add_bullet("Mobile app for patients (optional)")
add_bullet("Optional: real LLM integration in AI Practice Manager")

# ── 17. Project Statistics ───────────────────────────────────────────────
add_heading("17. Project Statistics", level=1)
add_kv_table([
    ("JavaScript LoC (server)", "~17,200"),
    ("HTML LoC (33 public pages)", "~7,400"),
    ("CSS LoC", "902"),
    ("i18n JSON (12 languages)", "906 lines combined"),
    ("Route modules", "17"),
    ("Database tables", "18"),
    ("Cron jobs", "8"),
    ("Languages supported", "12 (incl. RTL Arabic)"),
    ("QA checklist points", "187"),
    ("Documentation lines", "~600 (README + DEPLOYMENT + QA_CHECKLIST)"),
    ("Active development phases shipped", "16"),
])

# ── 18. Opportunities ────────────────────────────────────────────────────
add_heading("18. Opportunities", level=1)
add_bullet("Re-enable Stripe → activate revenue from day 1 of new ownership.")
add_bullet("Vertical specialization (physio, beauty, coaching) using existing data model + new UI labels.")
add_bullet("White-label / SaaS-for-SaaS — sell the platform to chains of practices.")
add_bullet("Geographic launch in MENA / SE Asia using existing AR/TR/ID/TH/ZH/HI catalogs.")
add_bullet("Add real LLM summary on top of the rule engine → premium tier upsell.")
add_bullet("Public-facing review aggregator / discovery page → SEO compound effect.")

# ── 19. Risks ────────────────────────────────────────────────────────────
add_heading("19. Risks", level=1)
add_kv_table([
    ("Market", "Doctolib, samedi, Jameda are entrenched in DACH; PraxisOnline24 has no marketplace effect."),
    ("Technical", "SQLite limits scale to ~200 practices; migration to PostgreSQL is documented but not done."),
    ("Operational", "No automated tests; production hardening (CSP/CSRF/2FA) outstanding."),
    ("Translation", "RU/ZH/HI/TH likely machine-translated; native review recommended."),
    ("Legal", "AGB/Datenschutz are templates; require lawyer review before live use under buyer brand."),
    ("Pre-revenue", "No customer signal; product-market fit is unproven."),
])

doc.add_page_break()

# ── 20. Market Potential ─────────────────────────────────────────────────
add_heading("20. Market Potential", level=1)
add_para(
    "Conservative bottoms-up: capture 100 paying practices at €59/month average → "
    "€5,900 MRR / €70,800 ARR within 12 months of focused execution.")
add_para(
    "Optimistic: capture 500 practices across DACH + 1–2 secondary language markets → "
    "€29,500 MRR / €354,000 ARR within 24 months.")
add_para("Comparable SaaS in the same niche (e.g., samedi, Jameda Pro) demonstrate that the "
         "willingness-to-pay exists and that retention is high once a practice integrates the tool "
         "into daily operations.")

# ── 21. Valuation ────────────────────────────────────────────────────────
add_heading("21. Valuation", level=1)
add_para("Methodology: replacement cost + comparable pre-revenue SaaS sales (Acquire.com, MicroAcquire). "
         "Full reasoning in Pricing_Strategy.md.", italic=True)
add_kv_table([
    ("Replacement cost (lower bound)", "€75,000  / ~$80,000"),
    ("Replacement cost (upper bound)", "€120,000 / ~$130,000"),
    ("Quick Sale (10–15% of repl.)", "$8,000 – $15,000"),
    ("Fair Market Value (25–35%)", "$18,000 – $32,000 — recommended list $24,500"),
    ("Optimistic (40–60%)", "$35,000 – $55,000"),
])

# ── 22. Sales Arguments ──────────────────────────────────────────────────
add_heading("22. Sales Arguments", level=1)
add_bullet("Production-deployed, live URL — verifiable in 1 minute.")
add_bullet("12-language coverage including RTL Arabic — unique in this price band.")
add_bullet("Privacy-by-design data model — no health data stored.")
add_bullet("8 cron-driven automations — operational scaffolding already done.")
add_bullet("AI Practice Manager — instantly upgradable to LLM-backed.")
add_bullet("Multi-tenant from day one — practice_id everywhere.")
add_bullet("Documented deployment via render.yaml — buyer redeploys in 5 minutes.")
add_bullet("Brand + domain + Cloudflare zone included.")
add_bullet("No framework lock-in — vanilla HTML/CSS/JS frontend.")
add_bullet("Free pre-acquisition demo guide + open access to read-only repo on signed NDA.")

doc.add_page_break()

# ── Closing ──────────────────────────────────────────────────────────────
add_heading("Contact & Next Steps", level=1)
add_para("Demo: live at https://www.praxisonline24.com")
add_para("Public listing copy: see Acquire_Listing.md")
add_para("Technical detail: see Technical_Due_Diligence.md")
add_para("Pricing reasoning: see Pricing_Strategy.md")
add_para("Buyer questions: see Buyer_FAQ.md")
add_para("Handover process: see Transfer_Checklist.md")
add_para("Demo walkthrough: see Demo_Guide.md")
add_para("")
add_para(
    "This report and its supporting documents are provided in good faith and reflect the actual state "
    "of the codebase, deployment, and assets at the time of writing. Any forward-looking statements "
    "(market potential, valuation) are reasoned estimates, not guarantees. Buyer is encouraged to "
    "perform independent due diligence.", italic=True, size=10)

doc.save(OUT)
print(f"Written: {OUT}")
